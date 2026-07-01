from pathlib import Path
from docx import Document

SOURCE = Path(r"C:\Users\user\Dropbox\EMS\院前EKG_APP_Codex執行提示版.docx")
OUT = Path(r"C:\Users\user\Documents\院前EKG\docx_extracted.txt")


def main():
    doc = Document(SOURCE)
    lines = [f"# {SOURCE.name}", ""]

    lines.append("## Paragraphs")
    for index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            lines.append(f"{index}. {text}")

    lines.append("")
    lines.append("## Tables")
    for table_index, table in enumerate(doc.tables, start=1):
        lines.append(f"### Table {table_index}")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")

    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(OUT)


if __name__ == "__main__":
    main()
